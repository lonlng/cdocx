#!/usr/bin/env python3
"""Generate a test DOCX with bookmarks for TemplateEngine validation."""

from docx import Document
from docx.shared import Pt

doc = Document()

# Add paragraph with bookmarks
p = doc.add_paragraph()
p.add_run("Company: ")
run = p.add_run("PLACEHOLDER_COMPANY")
run.bold = True

# Add bookmark manually via XML
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# This is a simplified approach - for real testing we'll use cdocx BookmarkInserter
# But for now we create a simple docx with placeholder text
doc.add_paragraph("Report Title: {{report_title}}")
doc.add_paragraph("Date: {{report_date}}")

p = doc.add_paragraph()
p.add_run("Signature: ")
run = p.add_run("SIGNATURE_HERE")
run.italic = True

doc.save("test_template.docx")
print("Created test_template.docx")
