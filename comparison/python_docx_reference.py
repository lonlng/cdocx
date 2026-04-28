#!/usr/bin/env python3
"""
Generate reference DOCX files using python-docx for comparison with cdocx output.
These files demonstrate common features and serve as a baseline for XML quality analysis.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = "python_docx_output"
os.makedirs(OUT_DIR, exist_ok=True)


def create_basic_document():
    """Simple document with paragraphs, formatting, and a table."""
    doc = Document()

    # Title
    title = doc.add_heading("Test Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Normal paragraph
    p1 = doc.add_paragraph("This is a normal paragraph with ")
    p1.add_run("bold text").bold = True
    p1.add_run(" and ")
    p1.add_run("italic text").italic = True
    p1.add_run(".")

    # Styled paragraph
    p2 = doc.add_paragraph()
    run = p2.add_run("Styled run with font size and color")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Table
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell {i},{j}"

    # Numbered list
    doc.add_paragraph("First item", style="List Number")
    doc.add_paragraph("Second item", style="List Number")

    # Bullet list
    doc.add_paragraph("Bullet one", style="List Bullet")
    doc.add_paragraph("Bullet two", style="List Bullet")

    doc.save(f"{OUT_DIR}/01_basic_document.docx")
    print(f"Created {OUT_DIR}/01_basic_document.docx")


def create_bookmark_document():
    """Document with bookmarks around text."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()
    doc.add_paragraph("Document with bookmarks")

    p = doc.add_paragraph()
    p.add_run("Before bookmark ")

    # Add bookmark via low-level XML using proper namespace
    run = p.add_run("bookmarked text")
    nsmap = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    bookmark_start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="0" w:name="MY_BOOKMARK"/>'
    )
    bookmark_end = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="0"/>'
    )
    run._element.insert(0, bookmark_start)
    run._element.append(bookmark_end)

    p.add_run(" after bookmark.")

    doc.save(f"{OUT_DIR}/02_bookmark_document.docx")
    print(f"Created {OUT_DIR}/02_bookmark_document.docx")


def create_image_document():
    """Document with an image placeholder (no actual image file)."""
    doc = Document()
    doc.add_paragraph("Document with image")
    # python-docx requires an image file; skip for this comparison
    doc.save(f"{OUT_DIR}/03_image_placeholder.docx")
    print(f"Created {OUT_DIR}/03_image_placeholder.docx")


def create_empty_document():
    """Truly empty document."""
    doc = Document()
    doc.save(f"{OUT_DIR}/04_empty.docx")
    print(f"Created {OUT_DIR}/04_empty.docx")


def create_complex_formatting():
    """Document with advanced formatting."""
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("Complex formatting test")
    run.bold = True
    run.italic = True
    run.underline = True
    run.font.strike = True
    run.font.superscript = True
    run.font.size = Pt(18)

    # Paragraph with alignment and spacing
    p2 = doc.add_paragraph("Aligned paragraph")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.5

    # Shading (via low-level XML in python-docx)
    p3 = doc.add_paragraph("Shaded paragraph")
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    pPr = p3._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="DDDDDD" w:val="clear"/>')
    pPr.append(shd)

    doc.save(f"{OUT_DIR}/05_complex_formatting.docx")
    print(f"Created {OUT_DIR}/05_complex_formatting.docx")


def create_table_document():
    """Document with merged cells."""
    doc = Document()
    doc.add_paragraph("Table with merged cells")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Merge first row cells
    cell_a = table.rows[0].cells[0]
    cell_b = table.rows[0].cells[1]
    cell_c = table.rows[0].cells[2]
    cell_a.merge(cell_b)
    cell_a.merge(cell_c)
    cell_a.text = "Merged header"

    for i in range(3):
        for j in range(3):
            if not table.rows[i].cells[j].text:
                table.rows[i].cells[j].text = f"R{i}C{j}"

    doc.save(f"{OUT_DIR}/06_table_merge.docx")
    print(f"Created {OUT_DIR}/06_table_merge.docx")


def analyze_docx(path, label):
    """Print structural analysis of a DOCX file."""
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET

    print(f"\n{'='*60}")
    print(f"Analysis: {label}")
    print(f"{'='*60}")

    size = os.path.getsize(path)
    print(f"File size: {size} bytes")

    with ZipFile(path, "r") as zf:
        print(f"\nZIP contents ({len(zf.namelist())} entries):")
        for name in zf.namelist():
            info = zf.getinfo(name)
            print(f"  {name:40s} {info.file_size:8d} bytes ({info.compress_size:8d} compressed)")

        # Parse document.xml
        with zf.open("word/document.xml") as f:
            tree = ET.parse(f)
            root = tree.getroot()

            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            paragraphs = root.findall(".//w:p", ns)
            runs = root.findall(".//w:r", ns)
            tables = root.findall(".//w:tbl", ns)
            bookmarks_start = root.findall(".//w:bookmarkStart", ns)
            bookmarks_end = root.findall(".//w:bookmarkEnd", ns)
            text_nodes = root.findall(".//w:t", ns)
            rpr_nodes = root.findall(".//w:rPr", ns)
            ppr_nodes = root.findall(".//w:pPr", ns)

            print(f"\nXML structure:")
            print(f"  Paragraphs:        {len(paragraphs)}")
            print(f"  Runs:              {len(runs)}")
            print(f"  Tables:            {len(tables)}")
            print(f"  Bookmark starts:   {len(bookmarks_start)}")
            print(f"  Bookmark ends:     {len(bookmarks_end)}")
            print(f"  Text nodes:        {len(text_nodes)}")
            print(f"  Run properties:    {len(rpr_nodes)}")
            print(f"  Para properties:   {len(ppr_nodes)}")


if __name__ == "__main__":
    create_basic_document()
    create_bookmark_document()
    create_image_document()
    create_empty_document()
    create_complex_formatting()
    create_table_document()

    print("\n" + "=" * 60)
    print("PYTHON-DOCX REFERENCE ANALYSIS")
    print("=" * 60)
    for f in sorted(os.listdir(OUT_DIR)):
        if f.endswith(".docx"):
            analyze_docx(f"{OUT_DIR}/{f}", f"python-docx: {f}")
