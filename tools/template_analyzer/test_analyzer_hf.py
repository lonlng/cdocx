#!/usr/bin/env python3
"""
Test Template Analyzer Header/Footer Detection
"""

import os
import subprocess
import sys
from docx import Document

ANALYZER = os.path.abspath(os.path.join(os.path.dirname(__file__), "analyze_template.py"))


def test_analyzer_detects_hf_placeholders():
    """Test that the analyzer detects placeholders in headers and footers."""
    print("=" * 60)
    print("TEST: Template Analyzer Header/Footer Detection")
    print("=" * 60)

    template_path = os.path.abspath("test_analyzer_hf.docx")
    output_cpp = os.path.abspath("test_analyzer_hf_output.cpp")

    # Create docx with placeholders in body, header, footer
    doc = Document()
    doc.add_paragraph("Company: {{company}}")
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = "Report: {{report_title}}"
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = "Page {{page_num}} of {{total_pages}}"
    doc.save(template_path)

    # Run analyzer
    result = subprocess.run(
        [sys.executable, ANALYZER, template_path, '-o', output_cpp],
        capture_output=True, text=True
    )
    print(f"Analyzer exit: {result.returncode}")
    print(f"Output: {result.stdout.strip()}")
    if result.stderr:
        print(f"Stderr: {result.stderr.strip()}")

    success = False
    if os.path.exists(output_cpp):
        with open(output_cpp, 'r', encoding='utf-8') as f:
            content = f.read()
        has_company = 'company' in content
        has_report = 'report_title' in content
        has_page = 'page_num' in content
        has_total = 'total_pages' in content
        success = has_company and has_report and has_page and has_total
        print(f"Keys found: company={has_company}, report_title={has_report}, page_num={has_page}, total_pages={has_total}")
        print(f"ASSERT: {'PASS' if success else 'FAIL'}")

    # Cleanup
    for f in [template_path, output_cpp]:
        if os.path.exists(f):
            os.remove(f)

    return success


def test_analyzer_detects_hf_bookmarks():
    """Test that the analyzer detects bookmarks in headers and footers."""
    print("\n" + "=" * 60)
    print("TEST: Template Analyzer Header/Footer Bookmark Detection")
    print("=" * 60)

    template_path = os.path.abspath("test_analyzer_hf_bm.docx")
    output_cpp = os.path.abspath("test_analyzer_hf_bm_output.cpp")

    doc = Document()
    doc.add_paragraph("Body text")
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = hp.add_run("HEADER_BM_TEXT")

    from docx.oxml.ns import qn
    p = hp._p
    r = run._r
    bm_start = p.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): '0',
        qn('w:name'): 'HEADER_BOOKMARK'
    })
    bm_end = p.makeelement(qn('w:bookmarkEnd'), {qn('w:id'): '0'})
    p.insert(p.index(r), bm_start)
    p.insert(p.index(r) + 1, bm_end)

    doc.save(template_path)

    result = subprocess.run(
        [sys.executable, ANALYZER, template_path, '-o', output_cpp],
        capture_output=True, text=True
    )
    print(f"Analyzer exit: {result.returncode}")
    print(f"Output: {result.stdout.strip()}")
    if result.stderr:
        print(f"Stderr: {result.stderr.strip()}")

    success = False
    if os.path.exists(output_cpp):
        with open(output_cpp, 'r', encoding='utf-8') as f:
            content = f.read()
        has_bm = 'HEADER_BOOKMARK' in content
        success = has_bm
        print(f"Bookmark found: {has_bm}")
        print(f"ASSERT: {'PASS' if success else 'FAIL'}")

    for f in [template_path, output_cpp]:
        if os.path.exists(f):
            os.remove(f)

    return success


def main():
    results = []
    results.append(("Analyzer HF Placeholders", test_analyzer_detects_hf_placeholders()))
    results.append(("Analyzer HF Bookmarks", test_analyzer_detects_hf_bookmarks()))

    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    passed = sum(1 for _, r in results if r)
    for name, result in results:
        print(f"  [{'PASS' if result else 'FAIL'}] {name}")
    print(f"\nTotal: {passed}/{len(results)} tests passed")
    return passed == len(results)


if __name__ == "__main__":
    ok = main()
    sys.exit(0 if ok else 1)
